import openpyxl, os
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt
from funktions import usr

def exceltoword(xlfiles, resultdir):
    document = Document()
    '''styles = document.styles
    for s in styles:
        document.add_paragraph(s.name)
    document.add_page_break()'''

    # take a list of xlsx files as input
    for file in xlfiles:
        wb = openpyxl.load_workbook(file)
        for ws in wb.get_sheet_names():
            sht = wb.get_sheet_by_name(ws)
            #document.add_heading(ws + ' ' + str(sht['A1'].value),level=1)
            if ws[-5:] == 'Chart':
                pass
            elif ws == 'Drug_Alcohol':
                cells = sht['A1':'M' + str(sht.max_row)]  # cells is a tuple
                print(len(cells))
                document.add_heading(ws + ' ' + str(sht['A1'].value), level=1)
                tbl = document.add_table(rows=len(cells), cols=13, style=document.styles['Table Grid'])
            elif ws[-4:] == 'F_II':
                cells = sht['A1':'F' + str(sht.max_row)]  # cells is a tuple
                print(len(cells))
                document.add_heading(ws + ' ' + str(sht['A1'].value), level=1)
                tbl = document.add_table(rows=len(cells), cols=6, style=document.styles['Table Grid'])
            elif ws[-12:] == 'Vehicle Type':
                cells = sht['A1':'H' + str(sht.max_row)]  # cells is a tuple
                print(len(cells))
                document.add_heading(ws + ' ' + str(sht['A1'].value), level=1)
                tbl = document.add_table(rows=len(cells), cols=8, style=document.styles['Table Grid'])
            else:
                cells =  sht['A1':'G'+str(sht.max_row)] # cells is a tuple
                print(len(cells))
                document.add_heading(ws + ' ' + str(sht['A1'].value), level=1)
                tbl = document.add_table(rows=len(cells),cols=7,style= document.styles['Table Grid'])
                # tbl.rows[1].alignment = WD_ALIGNMENT_CENTER
            #print(tbl.style.font.size)
            #tbl.style.font.size = Pt(12)
            #tbl.style = 'Grid Table4-Accent1'
            for index, row in enumerate(cells):
                #print(index, row)
                append_row = [cell.value for cell in row]
                #print(append_row)
                hdr_cells = tbl.rows[index].cells
                #print(hdr_cells)
                for subindex, item in enumerate(append_row):
                    # print(subindex, item)
                    hdr_cells[subindex].text = str(item)
            #print(type(hdr_cells))
            if ws[-4:] == 'F_II':
                n=1
                while n < len(cells):
                    a = tbl.cell(n,0)
                    b = tbl.cell(n+1,0)
                    b.text = ''
                    a.merge(b)
                    #print(a.text)
                    '''a.text.replace('None', ' ')
                    print(a.text)'''
                    n += 2

            '''if ws == 'Drug_Alcohol':
                i=1
                print(len(tbl.columns))
                while i < len(tbl.columns):
                    print(tbl.cell(0,i).text)
                    print(tbl.cell(0,i+1).text)
                    namecell = tbl.cell(0,i)
                    blankcell = tbl.cell(0,i+1)
                    blankcell.text = ''
                    print(blankcell.text)
                    namecell.merge(blankcell)  # does it merge differently left to right?
                    print(namecell.text)
                    i += 2'''
            tbl.autofit = True

            document.add_page_break()
            #frst_row = []
            #for i in range(sht.max_row):
            #for row in sht.iter_rows(max_row=sht.max_row,max_col=7):


    document.save(os.path.join(resultdir,'CrashReport.docx'))
    return os.path.join(resultdir,'CrashReport.docx')
if __name__ == "__main__":
    xlfiles = ['C:/Users/' + usr + '/Desktop/Top_Locations.xlsx',
               'C:/Users/' + usr + '/Desktop/Tables.xlsx',
               'C:/Users/' + usr + "/Desktop/Alcohol_Drug.xlsx",
               'C:/Users/' + usr + "/Desktop/CrashReportTableYear.xlsx",
               'C:/Users/' + usr + '/Desktop/Seatbelt.xlsx',
               'C:/Users/' + usr + '/Desktop/Overall_F_II.xlsx',
               'C:/Users/' + usr + '/Desktop/Tables_F_II.xlsx',
               'C:/Users/' + usr + "/Desktop/Drug_Alcohol_F_II.xlsx"

    ]

    exceltoword(xlfiles,"C:/Users/" + usr + "/Desktop")
