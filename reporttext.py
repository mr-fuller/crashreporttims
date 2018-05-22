from docx import Document
import os
from funktions import usr, report_yrs

def execsummary(resultdir):
    yrs = '2012-2014'  # import this from somewhere
    dcmnt = Document()
    dcmnt.add_paragraph('EXECUTIVE SUMMARY')
    intro_head = dcmnt.add_paragraph()
    intro_head.add_run('Introduction').bold = True
    intro = ('One of the tasks of the Toledo Metropolitan Area Council of Governments (TMACOG) as the Metropolitan Planning'
         ' Organization (MPO) for Lucas, Wood, and Monroe counties is to evaluate systems analyses and operations to identify'
         ' transportation needs of the region. A key element of the systems analyses is to identify locations on the roadway'
         ' network that are demonstrating higher than normal crash occurrences. To address the need to evaluate high crash '
         'locations, in August 2012, TMACOG started an update of the Safety Locations and Measures Report (2007-2009 Crash '
         'Data). Every state is required to develop a safety plan based on crash data to identify the leading causes of '
         'incapacitating injuries and deaths on public roadways. The intent of this Safety Locations and Measures Report '
         '(' + yrs + ' Crash Data) is to provide a generalized assessment on the level of roadway safety on local '
         'federal-aid-eligible route system. This report will use the shorthand name Safety Report ' + yrs[0:5]+ yrs[-2:]+'.')
    dcmnt.add_paragraph(intro)
    prps_head = dcmnt.add_paragraph()
    prps_head.add_run('Purpose').bold = True
    prps = ('Federal, state, and local planning authorities utilize crash records and incident locations to facilitate '
            'planning, prioritizing, monitoring, and analyzing the need for transportation improvements and the '
            'programming of projects. This crash data and analysis provides guidance on how best to appropriate '
            'limited fiscal resources available for capital improvements that will best improve roadway safety and '
            'get the most benefit for the dollar. The crash statistics/analyses are beneficial to local governmental '
            'agencies/jurisdictions engaged in transportation planning, highway safety and transportation '
            'engineering since the findings can be used to inform and encourage local jurisdictions to pursue state '
            'and federal safety funds for projects.')
    dcmnt.add_paragraph(prps)
    dcmnt.add_paragraph('The key objectives of the Safety Locations Report (' + yrs + ' Crash Data) are:')
    dcmnt.add_paragraph('Assess countywide overall crash statistic characteristics;', style='List Number')
    dcmnt.add_paragraph('Review fatal and incapacitating injury crash characteristics;', style='List Number')
    dcmnt.add_paragraph('Summarize ODOT identified high crash locations on non-freeway, state and U.S. routes;', style='List Number')
    dcmnt.add_paragraph('Define high crash criteria, identify high crash locations and rank these locations;', style='List Number')
    dcmnt.add_paragraph('Identify and review top 50 high crash intersections in the TMACOG planning area', style='List Number 2')
    dcmnt.add_paragraph('Identify and review top 50 high crash segments in the TMACOG planning area and', style='List Number 2')
    dcmnt.add_paragraph('Review current education measures and potential needs.', style='List Number')
    p = dcmnt.add_paragraph()
    p.add_run('Process').bold = True
    dcmnt.add_paragraph('ODOT\'s Highway Safety GIS Crash Analysis Tool (GCAT) web-based database and from the '
            'Michigan crash data site www.michigantrafficcrashfacts.org (MTCF). TMACOG utilized the GCAT '
            'program and the MTCF site to download crash databases for the years ' +yrs[0:4] +  ' to ' +yrs[-4:] +'. This data was '
            'used to perform the countywide crash statistics analyses as well as to identify high crash intersections '
            'and sections on the local federal-aid-eligible routes in Lucas, Wood and Monroe counties. The crash '
            'data was used to implement a ranking formula that was based on three key criteria to determine the '
            'highest top crash intersections and top segments. The three criteria used in the ranking formula '
            'involved:')
    criteria1 = dcmnt.add_paragraph(style='List Paragraph')
    criteria1.add_run("Crash Frequency").bold = True
    criteria1.add_run(' – In order to qualify as a high crash location to be evaluated in the ranking '
                        'process, intersections needed 10 or more crashes and sections needed 20 or more crashes in '  # add threshold variables here
                        'the three-year period.')
    criteria2 = dcmnt.add_paragraph(style='List Paragraph')
    criteria2.add_run("Crash Rate").bold = True
    criteria2.add_run(' – The crash rate for an intersection is based on the number of crashes per million '
                      'entering vehicles (MEV) to the intersection; and section rates are based on the number of '
                      'crashes per million vehicle miles traveled (MVMT).')
    criteria3 = dcmnt.add_paragraph(style= 'List Paragraph')
    criteria3.add_run('Equivalent Property Damage Only (EPDO)').bold = True
    criteria3.add_run(' – The EPDO is a composite measure of the crash severity of a location.')
    dcmnt.add_paragraph('Once the top intersections and segments on the local federal-aid-eligible routes were identified based ' 
                        'on their ranking formula scores, a general overview of each location was conducted including a '
                        'summary of key crash data. The local jurisdiction can then decide if they would like to pursue safety '
                        'funding for improving the identified high crash location by having a detailed engineering safety study '
                        'conducted with the current crash data available and apply to the Safety Funding Program for '
                        'state/federal funds. If the jurisdiction is interested the System Performance and Monitoring (SPAM) '
                        'committee will take a more in depth look at any location studied in this report.')
    dcmnt.add_paragraph('The source of data for the high crash locations on the state routes (SR), U.S. routes (non-freeway) '
                        'and interstates was the 2016 Safety Analyst Study Locations. This data was sorted for Lucas and '
                        'Wood counties and utilized to identify the top 20 crash listings that ODOT is studying this fiscal '
                        'year. Where these routes pass through incorporated municipalities, the local community is '
                        'responsible for the roadway facility. Given this, any of the locations on this listing that are located in '
                        'a municipality can be used to determine if a roadway is eligible to pursue safety funding for '
                        'improving the identified high crash location.')
    fndng = dcmnt.add_paragraph()
    fndng.add_run('Findings').bold = True
    dcmnt.add_paragraph('There were three main crash data sets that were reviewed in this report. The key findings of '
                        'study for these two data categories are outlined below.')
    dcmnt.add_paragraph('Countywide Crash Data Statistics')
    dcmnt.add_paragraph('Local Federal-Aid-Eligible Routes',style='List Paragraph')
    dcmnt.add_paragraph('Top 50 ranked high crash intersections', style='List Paragraph')
    dcmnt.add_paragraph('Top 50 ranked high crash segments', style='List Paragraph')
    dcmnt.add_paragraph('ODOT District 2 Safety Work Program', style='List Paragraph')
    stats_head = dcmnt.add_paragraph()
    stats_head.add_run('Countywide Crash Data Statistics').bold = True
    stats_head.add_run(' - There were a total of 51,298 crash records within the data set, '
                       'of which 38,930 were in Lucas County, 10,065 in Wood County and 2,845 in Monroe County. These '
                       'statistics were evaluated for all of the crashes, as well as a subset of analyses on just those crashes '
                       'that had a fatality or incapacitating injury for several select categories. Following are highlights from '
                       'the crash data records:')
    hghlght= dcmnt.add_paragraph()
    hghlght.add_run('Highlights of countywide statistics for ALL crashes:').underline = True
    dcmnt.add_paragraph('Overall crashes by year in all three counties indicated 17,477 in 2012; 16,744 in 2013; and'
'17,619 in 2014.',style='List Bullet')
    dcmnt.add_paragraph('The three most common types of crashes in Lucas County were rear-end, angle, and '
                        'sideswipe passing; in Wood County they were rear-end, fixed object, and angle; and in '
                        'Monroe County they were single motor vehicle, rear end, and angle.',style='List Bullet')
    dcmnt.add_paragraph('The two most common crash contributing factors listed for both Lucas and Wood counties '
                        'were following too closely/assured clear distance ahead and failure to control. No contributing ' 
                        'factor data was available for Monroe County.',style='List Bullet')
    dcmnt.add_paragraph('In Lucas and Wood counties, the day of the week on which the highest frequency of crashes '
                        'occurred was on Friday. In Monroe County the day of the week on which the highest '
                        'frequency of crashes occurred was on Thursday.',style='List Bullet')
    dcmnt.add_paragraph('The highest frequency of crashes in Lucas, Wood, and Monroe counties based on the time of '
                        'day occurred in the p.m. peak from 3-6 p.m.',style='List Bullet')
    hghlght_f_ii = dcmnt.add_paragraph()
    hghlght_f_ii.add_run('Highlights of countywide statistics for ONLY crashes with fatalities and/or incapacitating injuries:').underline = True
    dcmnt.add_paragraph('In the 2009 to 2011 period there were 177 crashes (106 in Lucas, 49 in Wood, and 22 in '
                        'Monroe) that involved a fatality.',style='List Bullet')
    dcmnt.add_paragraph('In the 2009 to 2011 period there were 1,571 crashes (1,126 in Lucas, 380 in Wood, and 65 in '
                        'Monroe) that involved an incapacitating injury.',style='List Bullet')
    dcmnt.add_paragraph('In all three counties the most common types of fatal/ incapacitating injury crashes are angle, '
                        'rear-end, and fixed object.',style='List Bullet')
    dcmnt.add_paragraph('The most common contributing factors of the fatal/ incapacitating injury crashes in Lucas and '
                        'Wood counties were failure to control, failure to yield, and following too closely/assured clear '
                        'distance ahead. There is no contributing factor data for Monroe County.',style='List Bullet')
    dcmnt.add_paragraph('A total of 66 fatal crashes occurred (42 in Lucas, 11 in Wood, and 13 in Monroe) involving '
                        'alcohol as a contributing factor, which accounts for 37% of the total 177 fatal crashes.',style='List Bullet')
    dcmnt.add_paragraph('The highest category of type of location for fatal/incapacitating injury crashes in Lucas '
                        'County was intersection locations and in Wood and Monroe counties it was non-intersection '
                        'locations.',style='List Bullet')
    dcmnt.add_paragraph('The time period when the highest number of fatal crashes occur in Lucas County is 10 p.m.- '
                        '11 p.m.; and in Wood County they occur during the 3 p.m.-4 p.m. hour. Fatal crashes in '
                        'Monroe County occur mostly in the afternoon hours.',style='List Bullet')
    last_bullet =dcmnt.add_paragraph('Additional categories are displayed in ',style='List Bullet')
    last_bullet.add_run('Section 5.3').italic = True
    last_bullet.add_run(' of this safety report')

    fedAidRoutes = dcmnt.add_paragraph()
    fedAidRoutes.add_run('Local Federal-Aid-Eligible Routes (excluding state and U.S. routes)').bold = True
    fedAidRoutes.add_run(' – The primary goal of the Safety Report ' + yrs[0:5] + yrs[-2:] + ' is to identify the top high crash '
                         'intersections and segments within the TMACOG region of Ohio and Michigan on the local '
                         'federal-aid-eligible routes. This was accomplished by applying a crash ranking to those '
                         'locations that met a minimum frequency threshold. The locations are listed below on the next '
                         'two tables.').bold = False

    # Insert tables

    #After the tables
    dcmnt.add_paragraph('Each of these locations, as well as any other location that was studied in this report, '
                        'can be studied further by the SPAM committee. An initiative to pursue safety funding for '
                        'potential safety project locations is the responsibility of the local jurisdictions.')
    odotswp_run = dcmnt.add_paragraph()
    odotswp_run.add_run('ODOT District 2 Safety Work Program').bold = True
    odotswp_run.add_run(' – There were a total of 20 high crash corridors located in the two counties that are typically '
                        'less than a half a mile in length with a high number of crashes in a three-year period. The '
                        'table following shows the ODOT safety analyst study locations that were found in Lucas and '
                        'Wood counties.')





    dcmnt.save(os.path.join(resultdir,'ReportText.docx'))
    return os.path.join(resultdir,'ReportText.docx')
if __name__ == "__main__":
    resultdir = 'C:/Users/' + usr +'/Desktop'
    execsummary(resultdir)